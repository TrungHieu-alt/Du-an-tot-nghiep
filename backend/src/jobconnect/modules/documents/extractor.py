"""Text extraction from uploaded document files.

Supports PDF (via pypdf) and DOCX (via python-docx).
Any extraction failure returns empty string so the worker marks the job failed.
"""
from __future__ import annotations

import io
from typing import BinaryIO


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"


def extract_text(stream: BinaryIO, mime_type: str) -> str:
    """Extract plain text from a file stream. Returns empty string on failure."""
    data = stream.read()
    if mime_type == "application/pdf":
        return _extract_pdf(data)
    if mime_type == DOCX_MIME:
        return _extract_docx(data)
    # Legacy .doc (binary MS Word 97-2003) not supported in MVP.
    return ""


def _extract_pdf(data: bytes) -> str:
    try:
        import pypdf  # type: ignore[import]

        reader = pypdf.PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            parts.append(text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # type: ignore[import]

        document = docx.Document(io.BytesIO(data))
        parts: list[str] = []
        # Paragraphs
        for para in document.paragraphs:
            text = (para.text or "").strip()
            if text:
                parts.append(text)
        # Tables (common in CVs/JDs)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""
